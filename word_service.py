import re
from tkinter import messagebox

from docx import Document
from docx.shared import Pt


def get_default_docx_filename(selected_topic):
    topic_for_name = selected_topic.strip()
    if not topic_for_name:
        topic_for_name = "artykul"

    safe_name = topic_for_name.replace(" ", "_")
    safe_name = re.sub(r"[^a-zA-Z0-9_ąćęłńóśźżĄĆĘŁŃÓŚŹŻ-]", "", safe_name)
    if not safe_name:
        safe_name = "artykul"

    return f"{safe_name}.docx"


def save_to_word(article, file_name):
    if not article:
        messagebox.showwarning("Uwaga", "Najpierw wygeneruj artykuł")
        return

    document = Document()

    lines = article.splitlines()
    non_empty_lines = [line.strip() for line in lines if line.strip()]

    if not non_empty_lines:
        messagebox.showwarning("Uwaga", "Najpierw wygeneruj artykuł")
        return

    title_line = non_empty_lines[0]
    if title_line.lower().startswith("tytuł:"):
        title_text = title_line.split(":", 1)[1].strip()
    else:
        title_text = title_line

    title_heading = document.add_heading(title_text, level=1)
    title_heading.paragraph_format.space_after = Pt(10)

    section_headers = {
        "wstep": "Wstęp",
        "rozwiniecie": "Rozwinięcie",
        "podsumowanie": "Podsumowanie",
    }

    paragraph_lines = []

    def normalize_for_match(text):
        polish_chars = str.maketrans("ąćęłńóśźż", "acelnoszz")
        return text.strip().lower().rstrip(":").translate(polish_chars)

    def add_section_heading(text):
        heading = document.add_heading(text, level=2)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

    def flush_paragraph():
        if not paragraph_lines:
            return

        paragraph_text = " ".join(paragraph_lines).strip()
        if paragraph_text:
            paragraph = document.add_paragraph(paragraph_text)
            paragraph.paragraph_format.space_after = Pt(8)

        paragraph_lines.clear()

    for line in lines[1:]:
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        normalized_line = normalize_for_match(stripped)
        if normalized_line in section_headers:
            flush_paragraph()
            add_section_heading(section_headers[normalized_line])
        else:
            paragraph_lines.append(stripped)

    flush_paragraph()

    document.save(file_name)

    messagebox.showinfo("Sukces", "Plik zapisany pomyślnie")
