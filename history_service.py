from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document


HISTORY_FILE = Path(__file__).resolve().parent / "historia.docx"
SEPARATOR = "---"
HEADER_PATTERN = re.compile(r"^DATA:\s*(.*?)\s*\|\s*TYP:\s*(.*?)\s*\|\s*TEMAT:\s*(.*)$")


def _build_header(date_text: str, content_type: str, topic: str) -> str:
    return f"DATA: {date_text} | TYP: {content_type} | TEMAT: {topic}"


def _split_entry_text(entry_text: str) -> list[str]:
    if not entry_text:
        return []
    return entry_text.splitlines()


def _parse_entries() -> list[dict]:
    if not HISTORY_FILE.exists():
        return []

    document = Document(HISTORY_FILE)
    blocks: list[list[str]] = []
    current_block: list[str] = []

    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line == SEPARATOR:
            if current_block:
                blocks.append(current_block)
            current_block = []
            continue

        current_block.append(paragraph.text)

    if current_block:
        blocks.append(current_block)

    entries: list[dict] = []
    for idx, block in enumerate(blocks, start=1):
        meaningful = [line for line in block if line.strip()]
        if not meaningful:
            continue

        header_line = meaningful[0].strip()
        match = HEADER_PATTERN.match(header_line)

        if match:
            date_text = match.group(1).strip()
            content_type = match.group(2).strip()
            topic = match.group(3).strip()
            body_lines = meaningful[1:]
        else:
            date_text = ""
            content_type = ""
            topic = ""
            body_lines = meaningful

        text = "\n".join(body_lines).strip()
        entries.append(
            {
                "id": idx,
                "date": date_text,
                "topic": topic,
                "content_type": content_type,
                "text": text,
            }
        )

    return entries


def _write_entries(entries: list[dict]) -> None:
    document = Document()

    for entry in entries:
        document.add_paragraph(SEPARATOR)
        document.add_paragraph(_build_header(entry["date"], entry["content_type"], entry["topic"]))

        for line in _split_entry_text(entry["text"]):
            document.add_paragraph(line)

    document.save(HISTORY_FILE)


def save_article(topic: str, content_type: str, text: str) -> None:
    article_text = text.strip()
    if not article_text:
        return

    entries = _parse_entries()

    entries.append(
        {
            "id": len(entries) + 1,
            "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "topic": topic.strip(),
            "content_type": content_type.strip(),
            "text": article_text,
        }
    )
    _write_entries(entries)


def load_history() -> list[dict]:
    entries = _parse_entries()
    return list(reversed(entries))


def delete_article(article_id: int) -> None:
    entries = _parse_entries()
    filtered = [entry for entry in entries if entry["id"] != article_id]
    _write_entries(filtered)
